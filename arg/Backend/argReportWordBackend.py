#HEADER
#                   arg/Backend/argReportWordBackend.py
#               Automatic Report Generator (ARG) v. 1.0
#
# Copyright 2020 National Technology & Engineering Solutions of Sandia, LLC
# (NTESS). Under the terms of Contract DE-NA0003525 with NTESS, the U.S.
# Government retains certain rights in this software.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice,
#   this list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the copyright holder nor the names of its
#   contributors may be used to endorse or promote products derived from this
#   software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE
# ARE DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE
# LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR
# CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF
# SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS
# INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN
# CONTRACT, STRICT LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE)
# ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE
# POSSIBILITY OF SUCH DAMAGE.
#
# Questions? Visit gitlab.com/AutomaticReportGenerator/arg
#
#HEADER

import datetime
import os
import sys

import docx

from arg.Backend.argWordBackend import argWordBackend


class argReportWordBackend(argWordBackend):
    """A concrete class providing a Word backend to generate a report
    """

    def __init__(self, parameters):

        # Call superclass init
        super().__init__(parameters)

    def create_document_preamble(self, version=None):
        """Create document and its preamble given an Assembler version
        """

        # Prepare report prefix and output directory
        if not os.path.isdir(self.Parameters.OutputDir):
            # Try to create output directory if it does not exist yet
            print("[argReportWordBackend] Creating output directory {}".format(self.Parameters.OutputDir))
            try:
                os.makedirs(self.Parameters.OutputDir, 0o750)
            except OSError:
                print("** ERROR: could not create output directory. Exiting.")
                sys.exit(1)

        # Create Word output
        self.Report = docx.Document()

        # Title, organization, addres, authors, numbers, etc.
        self.Report.add_heading("{}\n{}\n{}".format(
            self.Parameters.Title,
            ", ".join(self.Parameters.Authors),
            datetime.date.today()), 0)

        # Include abstract if available
        if self.Parameters.Abstract:
            p = self.Report.add_paragraph()
            p.add_run("Abstract:").bold = True
            with open(os.path.join(self.Parameters.DataDir, self.Parameters.Abstract), 'r') as f:
                abstract_text = ''.join(f.readlines())
                self.Report.add_paragraph(abstract_text)

        # Close title page
        self.Report.add_page_break()

        # Add document timestamp
        self.add_document_provenance(version)

        # Append tables of contents, figures, tables
        self.add_document_tocs()
