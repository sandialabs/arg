# HEADER
#                      arg/Backend/argWordBackend.py
#               Automatic Report Generator (ARG) v. 1.0
#
# Copyright 2020 National Technology & Engineering Solutions of Sandia, LLC
# (NTESS). Under the terms of Contract DE-NA0003525 with NTESS, the U.S.
# Government retains certain rights in this software.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
#
# * Redistributions of source code must retain the above copyright notice,
#   this list of conditions and the following disclaimer.
#
# * Redistributions in binary form must reproduce the above copyright notice,
#   this list of conditions and the following disclaimer in the documentation
#   and/or other materials provided with the distribution.
#
# * Neither the name of the copyright holder nor the names of its
#   contributors may be used to endorse or promote products derived from this
#   software without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE
# ARE DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT OWNER OR CONTRIBUTORS BE
# LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR
# CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF
# SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS
# INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN
# CONTRACT, STRICT LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE)
# ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE
# POSSIBILITY OF SUCH DAMAGE.
#
# Questions? Visit gitlab.com/AutomaticReportGenerator/arg
#
# HEADER
from io import BytesIO
import numbers
import os
import platform
import re
import subprocess

import clr
import yaml
import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Mm, Cm, Inches, RGBColor, Pt
from pylatex import NoEscape

from arg.Common.argMultiFontStringHelper import argMultiFontStringHelper
from arg.Backend.argBackendBase import argBackendBase, ArgHTMLParser

osn = os.name

# Try importing Word Application in Python environment
try:
    clr.AddReference("System")
    clr.AddReference("Microsoft.Office.Interop")
    from Microsoft.Office.Interop import Word

# Find work-around
except:
    try:
        # Use Windows32 client on Windows only
        if osn == "nt":
            from win32com.client import DispatchEx
    except:
        pass

# Load supported colors
common_dir = os.path.dirname(os.path.realpath(__file__))
with open(os.path.join(common_dir, "../Common/argTypes.yml"),
          'r',
          encoding="utf-8") as t_file:
    supported_types = yaml.safe_load(t_file)

# Retrieve supported colors
colors = supported_types.get("FontColors")

# Retrieve supported fonts
fonts = supported_types.get("BackendTypes").get("Word").get("Fonts")


class argWordBackend(argBackendBase):
    """A concrete class providing a Word backend
    """

    # Map from greek letters to their respective unicodes
    GreekLetters = {
        "Alpha": u'\u0391',
        "Beta": u'\u0392',
        "Gamma": u'\u0393',
        "Delta": u'\u0394',
        "Epsilon": u'\u0395',
        "Zeta": u'\u0396',
        "Eta": u'\u0397',
        "Theta": u'\u0398',
        "Iota": u'\u0399',
        "Kappa": u'\u039A',
        "Lamda": u'\u039B',
        "Mu": u'\u039C',
        "Nu": u'\u039D',
        "Xi": u'\u039E',
        "Omicron": u'\u039F',
        "Pi": u'\u03A0',
        "Rho": u'\u03A1',
        "Sigma": u'\u03A3',
        "Tau": u'\u03A4',
        "Upsilon": u'\u03A5',
        "Phi": u'\u03A6',
        "Chi": u'\u03A7',
        "Psi": u'\u03A8',
        "Omega": u'\u03A9',
        "alpha": u'\u03B1',
        "beta": u'\u03B2',
        "gamma": u'\u03B3',
        "delta": u'\u03B4',
        "epsilon": u'\u03B5',
        "zeta": u'\u03B6',
        "eta": u'\u03B7',
        "theta": u'\u03B8',
        "iota": u'\u03B9',
        "kappa": u'\u03BA',
        "lamda": u'\u03BB',
        "mu": u'\u03BC',
        "nu": u'\u03BD',
        "xi": u'\u03BE',
        "omicron": u'\u03BF',
        "pi": u'\u03C0',
        "rho": u'\u03C1',
        "sigma": u'\u03C3',
        "tau": u'\u03C4',
        "upsilon": u'\u03C5',
        "phi": u'\u03C6',
        "chi": u'\u03C7',
        "psi": u'\u03C8',
        "omega": u'\u03C9'}

    def __init__(self, parameters=None):

        # Call superclass init
        super().__init__(parameters)

        # Define backend type
        self.Type = "Word"

        self.colors = None
        self.fonts = None
        self.app = "argWordBackend"
        self.osn = os.name

        # Try importing Word Application in Python environment
        try:
            clr.AddReference("System")
            clr.AddReference("Microsoft.Office.Interop")
            from Microsoft.Office.Interop import Word

        # Find work-around
        except:
            try:
                # Use Windows32 client on Windows only
                if self.osn == "nt":
                    from win32com.client import DispatchEx
            except:
                pass
        self.__init_fonts_colors()

    def __init_fonts_colors(self):
        # Load supported colors
        common_dir = os.path.dirname(os.path.realpath(__file__))
        with open(os.path.join(common_dir, "../Common/argTypes.yml"),
                  'r',
                  encoding="utf-8") as t_file:
            supported_types = yaml.safe_load(t_file)

        # Retrieve supported colors
        self.colors = supported_types.get("FontColors")

        # Retrieve supported fonts
        self.fonts = supported_types.get("BackendTypes").get("Word").get("Fonts")

    @staticmethod
    def generate_text(text, type="default"):
        """Add page break to the report
        """

        # Generate a backend dependent text
        return "{}".format(text)

    def generate_multi_font_string(self, multi_font_string: argMultiFontStringHelper, paragraph=None):
        """Either append Word runs to paragraph or return text string
        """

        # Append Word runs to paragraph when one was passed
        if paragraph:

            # Iterate over all items in string helper
            for string, font_bits, color, highlight_color in multi_font_string.iterator():
                if isinstance(font_bits, int):
                    # Map Greek letter to unicode when requested
                    if font_bits == 16:
                        string = self.GreekLetters.get(string, '')

                    # Add a run per string and set font per defined bits
                    run = paragraph.add_run(string)

                    if font_bits == 1:
                        run.font.italic = True
                    elif font_bits == 2:
                        run.font.bold = True
                    elif font_bits == 3:
                        run.font.underline = True
                    elif font_bits == 4:
                        run.font.name = self.fonts.get("typewriter", "default" if self.fonts.get("default") else None)
                    elif font_bits == 5:
                        run.font.strike = True
                    elif font_bits == 6:
                        run.font.subscript = True
                    elif font_bits == 8:
                        run.font.name = self.fonts.get("calligraphic", "default" if self.fonts.get("default") else None)
                    elif font_bits == 9:
                        run.font.size = Pt(14)
                        run.font.bold = True
                    elif font_bits == 10:
                        run.font.size = Pt(13)
                        run.font.bold = True
                    elif font_bits == 11:
                        run.font.size = Pt(11)
                        run.font.bold = True
                    elif font_bits == 12:
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.italic = True
                    elif font_bits == 13:
                        run.font.size = Pt(11)
                    elif font_bits == 14:
                        run.font.size = Pt(11)
                        run.font.italic = True
                    else:
                        run.font.name = self.fonts.get("default", None)

                elif isinstance(font_bits, list):
                    if 16 in font_bits:
                        string = self.GreekLetters.get(string, '')
                    run = paragraph.add_run(string)
                    run.font.italic = True if 1 in font_bits else False
                    run.font.bold = True if 2 in font_bits else False
                    run.font.underline = True if 3 in font_bits else False
                    run.font.strike = True if 5 in font_bits else False
                    run.font.subscript = True if 6 in font_bits else False
                    if 4 in font_bits:
                        run.font.name = self.fonts.get("typewriter", "default" if self.fonts.get("default") else None)
                    elif 8 in font_bits:
                        run.font.name = self.fonts.get("calligraphic", "default" if self.fonts.get("default") else None)
                    elif 9 in font_bits:
                        run.font.size = Pt(14)
                    elif 10 in font_bits:
                        run.font.size = Pt(13)
                    elif 11 in font_bits:
                        run.font.size = Pt(11)
                    elif 12 in font_bits:
                        run.font.size = Pt(11)
                    elif 13 in font_bits:
                        run.font.size = Pt(11)
                    elif 14 in font_bits:
                        run.font.size = Pt(11)
                    else:
                        run.font.name = self.fonts.get("default", None)

                elif isinstance(font_bits, dict):
                    run = paragraph.add_run(string)
                    if font_bits.get('font-size', None) is not None:
                        run.font.size = Pt(font_bits.get('font-size', 11))
                    if font_bits.get('font-family', None) is not None:
                        run.font.name = font_bits.get('font-family', None)

                if color:
                    if self.colors.get(color, None) is None:
                        rgb = color.split(',')
                        if len(rgb) > 2:
                            run.font.color.rgb = RGBColor(*[int(x) for x in rgb[:3]])
                    else:
                        rgb = self.colors.get(color).split(',')
                        if len(rgb) > 2:
                            run.font.color.rgb = RGBColor(*[int(x) for x in rgb[:3]])

                if highlight_color is not None:
                    color = self.get_closest_color(hex_str=highlight_color)
                    wd_color_mapping = {'AUTO': WD_COLOR_INDEX.AUTO, 'BLACK': WD_COLOR_INDEX.BLACK,
                                        'BLUE': WD_COLOR_INDEX.BLUE, 'BRIGHT_GREEN': WD_COLOR_INDEX.BRIGHT_GREEN,
                                        'DARK_BLUE': WD_COLOR_INDEX.DARK_BLUE, 'DARK_RED': WD_COLOR_INDEX.DARK_RED,
                                        'DARK_YELLOW': WD_COLOR_INDEX.DARK_YELLOW, 'GRAY_25': WD_COLOR_INDEX.GRAY_25,
                                        'GRAY_50': WD_COLOR_INDEX.GRAY_50, 'GREEN': WD_COLOR_INDEX.GREEN,
                                        'PINK': WD_COLOR_INDEX.PINK, 'RED': WD_COLOR_INDEX.RED,
                                        'TEAL': WD_COLOR_INDEX.TEAL, 'TURQUOISE': WD_COLOR_INDEX.TURQUOISE,
                                        'VIOLET': WD_COLOR_INDEX.VIOLET, 'WHITE': WD_COLOR_INDEX.WHITE,
                                        'YELLOW': WD_COLOR_INDEX.YELLOW}
                    run.font.highlight_color = wd_color_mapping.get(color.upper(), 'AUTO')

            # Nothing is returned in this case
            return

        # Return a plain string ignoring font specifications otherwise
        else:
            return ''.join([x[0] for x in multi_font_string.StringMap])

    def generate_matrix_string(self, matrix):
        """Generate Word string from matrix entries
        """

        # Bail out early if vector is empty or does not have correct type
        if not matrix or not isinstance(matrix, (list, tuple)):
            return ''

        # Return plain text string
        return ','.join(["[{}]".format(','.join([f"{x}" for x in v])) for v in matrix])

    @staticmethod
    def create_list(paragraph, list_type):
        """Create list
        """

        # Retrieve XML paragraph element
        p = paragraph._p

        # Retrieve paragraph properties
        pPr = p.get_or_add_pPr()

        # Create number properties element
        numPr = OxmlElement('w:numPr')

        # Create numId element to set bullet type
        numId = OxmlElement('w:numId')

        # Set list type/indentation
        numId.set(qn('w:val'), list_type)

        # Add bullet type to number properties list
        numPr.append(numId)

        # Add number properties to paragraph
        pPr.append(numPr)

    def add_list(self, item, num_lvl=1, number_items=False):
        """Add itemization or enumeration to the report
        """

        # Iterate over all elements of provided list
        for it in item['items']:
            # Add string elements
            if "string" in it.keys():
                self.Report.add_paragraph(it["string"])

            # Add numbered elements
            if "number" in it.keys():
                self.Report.add_paragraph(it["number"])

    def add_comment(self, comments_dict, key):
        """Add comment to the report from a dict as either text string
        or as sub-paragraph depending on number of dict values (1 or 2)
        """

        # Retrieve comment for given key and bail out if none found
        comment = comments_dict.get(str(key))
        if not comment:
            return False

        # Insert either sub-paragraph or string
        if len(comment) > 1:
            self.Report.add_paragraph(
                comment[0] + ":", "Subtitle")
            self.Report.add_paragraph(comment[1])
        else:
            self.Report.add_paragraph(comment[0])

        # Comment was found and inserted
        return True

    def add_paragraph(self, item, p=None):
        """Add paragraph to the report
        """
        if "hyperlink_path" in item:
            hyperlink_path = item.get('hyperlink_path', None)
            hyperlink_string = item.get('hyperlink_string', None)
            par = self.support_string(item=item) if "string" in item else self.Report.add_paragraph()
            _, par_hyp = self.add_hyperlink_to_paragraph(paragraph=par, url=hyperlink_path, text=hyperlink_string,
                                                         color='0000FF', underline=True)
            if "string_suffix" in item:
                string_suffix = item.get('string_suffix', None)
                par_str_suff = self.add_string_to_paragraph(paragraph=par_hyp, text=string_suffix)

        # Check whether a string or a file is to be included
        elif "string" in item:
            self.support_string(item=item)

        # Insert verbatim fragment
        elif "verbatim" in item:
            # Create paragraph in macro style
            p = self.Report.add_paragraph(item["verbatim"])
            p.style = self.Report.styles["macro"]

        # Insert contents of a file
        elif "include" in item and len(item["include"]) > 0:
            curr_item = item["include"][0]
            if "file" in curr_item:
                file_name = os.path.join(self.Parameters.DataDir, curr_item["file"])
                font = curr_item.get("font")
                size = curr_item.get("size")
                italic = curr_item.get("italic", False)
                underline = curr_item.get("underline", False)
                color = curr_item.get("color")
                alignment = curr_item.get("alignment")
                left_indent = curr_item.get("left_indent")
                right_indent = curr_item.get("right_indent")
                first_line_indent = curr_item.get("first_line_indent")
                space_before = curr_item.get("space_before")
                space_after = curr_item.get("space_after")
                line_spacing = curr_item.get("line_spacing")
                keep_with_next = curr_item.get("keep_with_next")
            else:
                file_name = os.path.join(self.Parameters.DataDir, item["include"])
                font, size, color = None, None, None
                italic, underline = False, False
                alignment = None
                left_indent, right_indent, first_line_indent = None, None, None
                space_before, space_after = None, None
                line_spacing = None
                keep_with_next = None

            with open(file_name, 'r') as tex_file:
                p = self.Report.add_paragraph()
                run = p.add_run()

                # Read each row and add it to Word output
                for line in tex_file:
                    run.text = run.text + line

                # Add font commands if provided
                if font == "typewriter":
                    run.font.name = self.fonts.get(font, "default" if self.fonts.get("default") else None)
                elif font == "calligraphic":
                    run.font.name = self.fonts.get(font, "default" if self.fonts.get("default") else None)
                if size:
                    run.font.size = Pt(size)
                if italic:
                    run.font.italic = italic
                if underline:
                    run.font.underline = underline
                if color:
                    rgb = color.split(',')
                    if len(rgb) > 2:
                        run.font.color.rgb = RGBColor(*[int(x) for x in rgb[:3]])
                if alignment:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if left_indent:
                    p.left_indent = Pt(left_indent)
                if right_indent:
                    run.right_indent = Pt(right_indent)
                if first_line_indent:
                    run.first_line_indent = Pt(first_line_indent)
                if space_before:
                    run.space_before = Pt(space_before)
                if space_after:
                    run.space_after = Pt(space_after)
                if line_spacing:
                    run.line_spacing = Pt(line_spacing)
                if keep_with_next:
                    run.keep_with_next = keep_with_next

    def add_subdivision(self, item, level):
        """Add specified subdivision to the report
        """

        # Retrieve title string
        title_string = item.get("title", '')

        # Handle title string according to its type
        if isinstance(title_string, str):
            # Directly insert base strings
            self.Report.add_heading(title_string, level)

        elif isinstance(title_string, argMultiFontStringHelper):
            # Handle multi-font string helpers
            self.generate_multi_font_string(
                title_string,
                self.Report.add_heading())

        # Add subdivision paragraph header if provided
        self.add_paragraph(item)

    def add_subtitle(self, item):
        """Add subsection to the report
        """

        # Call appropriate subdivision method
        self.add_subdivision(item, 3)

    def add_subsubsection(self, item, numbered=True):
        """Add subsection to the report
        """

        # Call appropriate subdivision method
        if self.check_orientation_change(item=item):
            self.change_orientation()
        self.add_subdivision(item, 4)

    def add_subsection(self, item, numbered=True):
        """Add subsection to the report
        """

        # Call appropriate subdivision method
        if self.check_orientation_change(item=item):
            self.change_orientation()
        self.add_subdivision(item, 3)

    def add_section(self, item, numbered=True):
        """Add section to the report
        """

        # Call appropriate subdivision method
        if self.check_orientation_change(item=item):
            self.change_orientation()
        self.add_subdivision(item, 2)

    def add_chapter(self, item, numbered=True):
        """Add chapter to the report
        """

        # Start with new page
        self.Report.add_page_break()

        # Call appropriate subdivision method
        if self.check_orientation_change(item=item):
            self.change_orientation()
        self.add_subdivision(item, 1)

    def add_page_break(self):
        """Add page break to the report
        """

        # Use eponymous command from Word backend
        self.Report.add_page_break()

    def add_heading(self, text, lvl=1, p=None):
        """Add heading to the report
        """

        # Use provided paragraph
        if p:
            p.text = text

        # Otherwise, use eponymous command from Word backend
        else:
            self.Report.add_heading(text, lvl)

    def add_table(self, header_list, body, caption_string, do_verbatim=False, _=None):
        """Add table to the report
        """

        # Bail out early if header is empty
        try:
            n_cols = len(header_list)
            if not n_cols:
                return
        except:
            return

        # Bail out early if unsupported body type
        if not body:
            return
        if type(body) == dict:
            is_dict = True
        elif type(body) == list:
            is_dict = False
        else:
            return

        # Create table with required number of columns
        table = self.Report.add_table(
            rows=1,
            cols=n_cols,
            style=self.Report.styles["Light List Accent 1"])
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create table header
        cells = table.rows[0].cells
        for c, header_string in zip(cells, header_list):
            # Fill header cell
            if isinstance(header_string, str):
                # Directly insert base strings
                c.text = header_string

            elif isinstance(header_string, argMultiFontStringHelper):
                # Handle multi-font string helpers
                self.generate_multi_font_string(
                    header_string,
                    c.paragraphs[0])

        # Iterate over of contents to create table body
        if is_dict:
            # Iterate over dictionary entries
            for k, v in sorted(body.items()):
                # Add a row of cells and fill those
                cells = table.add_row().cells

                # Make key and values lists if they are strings or numbers
                if isinstance(k, (str, numbers.Number)):
                    k = [k]
                if isinstance(v, (str, numbers.Number)):
                    v = [v]

                # Iterate to fill all cells
                for c, cell_string in zip(cells, list(k) + list(v)):
                    # Distinguish argMultiFontStrings from other types
                    if isinstance(cell_string, argMultiFontStringHelper):
                        # Handle multi-font string helpers
                        self.generate_multi_font_string(
                            "{}".format(cell_string),
                            c.paragraphs[0])
                    else:
                        # Directly insert base strings
                        c.text = "{}".format(cell_string)

                    # Use verbatim style when requested
                    if do_verbatim:
                        c.paragraphs[0].style = self.Report.styles["macro"]
        else:
            # Iterate over list entries
            for row in body:
                # Add a row of cells and fill those
                cells = table.add_row().cells
                for c, cell_string in zip(cells, row):
                    if isinstance(cell_string, str):
                        # Directly insert base strings
                        c.text = cell_string

                    elif isinstance(cell_string, argMultiFontStringHelper):
                        # Handle multi-font string helpers
                        self.generate_multi_font_string(
                            cell_string,
                            c.paragraphs[0])

                    # Use verbatim style when requested
                    if do_verbatim:
                        c.paragraphs[0].style = self.Report.styles["macro"]

        # Create table caption depending on its type
        if caption_string:
            if isinstance(caption_string, str):
                # Directly insert base strings
                caption = self.add_caption("Table", caption_string)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif isinstance(caption_string, argMultiFontStringHelper):
                # Handle multi-font string helpers
                caption = self.add_caption("Table", "")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Append Word runs created from multi-font string
                caption_string.execute_backend(caption)

            else:
                # Insert blank paragraph to ensure table is closed
                self.Report.add_paragraph()
        else:
            # Insert blank paragraph to ensure table is closed
            self.Report.add_paragraph()

    def change_orientation(self):
        """ Change orientation of the page from portrait to landscape or the other way around when called
        """
        current_section = self.Report.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = self.Report.add_section(WD_SECTION_START.NEW_PAGE)
        if self.Orientation == 'portrait':
            new_section.orientation = WD_ORIENTATION.LANDSCAPE
            self.Orientation = 'landscape'
        elif self.Orientation == 'landscape':
            new_section.orientation = WD_ORIENTATION.PORTRAIT
            self.Orientation = 'portrait'
        new_section.page_width = new_width
        new_section.page_height = new_height
        return new_section

    def check_orientation_change(self, item: dict) -> bool:
        """ Returns True if the page orientation needs to be changed. Returns False if no changes needed.
        """
        if item.get('orientation', None) is None and self.Orientation == 'portrait':
            return False
        elif item.get('orientation', None) is None and self.Orientation == 'landscape':
            return True
        elif item.get('orientation', None) == self.Orientation:
            return False
        elif item.get('orientation', None) != self.Orientation:
            return True

    def add_color_table(self, data: dict) -> None:
        """ Add colored table to the report
        """
        headers_list = data.get('headers', None)
        table_columns_num = len(headers_list)
        rows_list = data.get('rows', None)
        caption_string = data.get('caption', None)
        if headers_list is None or rows_list is None:
            raise Exception('No headers or rows given!')
        headers = self.parse_headers(headers=headers_list)
        rows = self.parse_rows(rows=rows_list, table_columns_num=table_columns_num)

        # Create table with required number of columns
        table = self.Report.add_table(rows=1, cols=table_columns_num)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create table header
        cells = table.rows[0].cells
        for c, header_string, font_color, alignment in \
                zip(cells, [header[0] for header in headers],
                    [self.parse_rgb_to_hex(rgb=header[2], appl=2) for header in headers],
                    [self.parse_cell_align(align=header[3]) for header in headers]):
            # Fill header cell
            if isinstance(header_string, str):
                # Directly insert base strings
                c.paragraphs[0].add_run(header_string).font.color.rgb = RGBColor.from_string(f"{font_color}")
                c.paragraphs[0].alignment = alignment


            elif isinstance(header_string, argMultiFontStringHelper):
                # Handle multi-font string helpers
                self.generate_multi_font_string(header_string, c.paragraphs[0])

        # Iterate over of contents to create table body

        # Iterate over list entries
        for row in rows:
            # Add a row of cells and fill those
            cells = table.add_row().cells
            for c, cell_string, font_color, alignment in \
                    zip(cells, [cell[0] for cell in row], [self.parse_rgb_to_hex(rgb=cell[2], appl=2) for cell in row],
                        [self.parse_cell_align(align=cell[3]) for cell in row]):
                if isinstance(cell_string, str):
                    # Directly insert base strings
                    c.paragraphs[0].add_run(cell_string).font.color.rgb = RGBColor.from_string(f"{font_color}")
                    c.paragraphs[0].alignment = alignment

                elif isinstance(cell_string, argMultiFontStringHelper):
                    # Handle multi-font string helpers
                    self.generate_multi_font_string(cell_string, c.paragraphs[0])

        # Apply colours to cells
        color_list = self.parse_colors(headers=headers, rows=rows)
        for row_num, row_color in enumerate(color_list):
            for cell_num, cell_color in enumerate(row_color):
                table.rows[row_num].cells[cell_num]._tc.get_or_add_tcPr().append(
                    parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{cell_color}"/>'))

        # Apply vertical alignment to cells:
        vertical_alignment = self.parse_vertical_alignment(headers=headers, rows=rows)
        for row_num, row_va in enumerate(vertical_alignment):
            for cell_num, cell_va in enumerate(row_va):
                table.rows[row_num].cells[cell_num]._tc.get_or_add_tcPr().append(
                    parse_xml(rf'<w:vAlign {nsdecls("w")} w:val="{cell_va}"/>'))

        # Create table caption depending on its type
        if caption_string is not None:
            if isinstance(caption_string, str):
                # Directly insert base strings
                caption = self.add_caption("Table", caption_string)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif isinstance(caption_string, argMultiFontStringHelper):
                # Handle multi-font string helpers
                caption = self.add_caption("Table", "")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Append Word runs created from multi-font string
                caption_string.execute_backend(caption)

            else:
                # Insert blank paragraph to ensure table is closed
                self.Report.add_paragraph()
        else:
            # Insert blank paragraph to ensure table is closed
            self.Report.add_paragraph()

    def inline_docx(self, data: dict) -> None:
        """ Add external docx document to the report
        """
        docx_path = data.get('docx_path', None)
        docx_path = os.path.join(self.Parameters.DataDir, docx_path)
        if docx_path is None:
            raise FileNotFoundError('Path to file NOT found!')
        elif not os.path.isfile(docx_path):
            raise FileNotFoundError('File does NOT exists!')
        else:
            # Mark in which paragraph there is an image
            inline_docx = docx.Document(docx_path)
            image_in_par = [True if 'graphicData' in par._p.xml else False for par in inline_docx.paragraphs]

            inline_docx = docx.Document(docx_path)
            # Setting counter coresponding to paragraphs
            counter = 0
            # Iterating over elements in document to inline
            for element in inline_docx.element.body:
                # If element is a paragraph, counter gets incremented
                if isinstance(element, docx.oxml.text.paragraph.CT_P):
                    # Eliminate use of style which does not exist
                    if element.style == 'Caption1':
                        element.style = 'Caption'
                    # Checking if paragraph had an image
                    if image_in_par[counter]:
                        # Taking an image from inline document and putting it to the last paragraph from current report
                        inl_docx = docx.Document(docx_path)
                        for shape in inl_docx.inline_shapes:
                            if shape.type == WD_INLINE_SHAPE.PICTURE:
                                # Make sure image isn't wider than 159mm, if it is, then set width do 159mm
                                if shape.width.mm > 159:
                                    ratio = shape.width.mm / 159
                                    pic_width = shape.width / ratio
                                else:
                                    pic_width = shape.width
                                # Taking image from inlining document
                                inline = shape._inline
                                rId = inline.xpath('./a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip/@r:embed')[0]
                                image_part = inl_docx._part.related_parts[rId]
                                image_bytes = image_part.blob
                                # write the image bytes to a file (or BytesIO stream)
                                image_stream = BytesIO(image_bytes)
                                last_par = self.Report.paragraphs[-1]
                                new_par = self.insert_paragraph_after(paragraph=last_par)
                                r = new_par.add_run()
                                # finally adding a picture with correct width
                                r.add_picture(image_stream, width=pic_width)
                        counter += 1
                    # Paragraph with no photo is simply added to the current report
                    else:
                        self.Report.element.body.append(element)
                        counter += 1
                # If element was not a paragraph is simply added to the current report
                else:
                    self.Report.element.body.append(element)

    def add_html(self, data: dict) -> None:
        """ Adds html content to Word document. """
        # Reading html data as a string
        html_str = data.get('html_string', None)
        # Parsing html data and get a html document in a form of list
        arg_html_parser = ArgHTMLParser()
        arg_html_parser.reset()
        html_list = arg_html_parser.get_mapped_html(html_str)
        # Parsing an html list into nested list over which is easier to iterate
        nested_list = self.nesting_html_list(html_list=html_list)
        # Parsing nested list to a list of mainly ARG Multi Font String Helper, ready to put in directly into document
        amfsh_list = self.map_nested_list_to_amfsh(nested_list=nested_list)
        # Final iteration over ARG Multi Font String Helper list
        for amfsh in amfsh_list:
            # Checking for tuple. Tuple instance is dedicated for any element except ordered and unordered lists
            if isinstance(amfsh, tuple):
                # amfsh as tuple: (alignment, argMultiFontStringHelper, margin)
                new_par = self.Report.add_paragraph()
                algn_map = {'CENTER': WD_ALIGN_PARAGRAPH.CENTER, 'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT, 'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY}
                new_par.alignment = algn_map.get(amfsh[0], WD_ALIGN_PARAGRAPH.LEFT)
                if amfsh[2] is not None and amfsh[2][0] == 'margin-right':
                    new_par.paragraph_format.right_indent = Pt(amfsh[2][1])
                elif amfsh[2] is not None and amfsh[2][0] == 'margin-left':
                    new_par.paragraph_format.left_indent = Pt(amfsh[2][1])
                self.generate_multi_font_string(multi_font_string=amfsh[1], paragraph=new_par)
            # Checking for list. List instance is dedicated for ordered and unordered lists (needs paragraph styling)
            elif isinstance(amfsh, list):
                # amfsh as list: [un/ordered list element, un/ordered list element, un/ordered list element, ...]
                for elem in amfsh:
                    # elem as list: [alignment, argMultiFontStringHelper, margin, list style]
                    new_par = self.Report.add_paragraph(style=elem[3])
                    algn_map = {'CENTER': WD_ALIGN_PARAGRAPH.CENTER, 'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                                'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT, 'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY}
                    new_par.alignment = algn_map.get(elem[0], WD_ALIGN_PARAGRAPH.LEFT)
                    if elem[2] is not None and elem[2][0] == 'margin-right':
                        new_par.paragraph_format.right_indent = Pt(elem[2][1])
                    elif elem[2] is not None and elem[2][0] == 'margin-left':
                        new_par.paragraph_format.left_indent = Pt(elem[2][1])
                    self.generate_multi_font_string(multi_font_string=elem[1], paragraph=new_par)

    @staticmethod
    def insert_paragraph_after(paragraph, text=None, style=None):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def add_hyperlink(self, data: dict) -> None:
        """ Add hyperlink to the report
        """
        path = data.get('path', None)
        hyperlink_string = data.get('hyperlink_string', None)
        par = self.Report.add_paragraph()
        self.add_hyperlink_to_paragraph(paragraph=par, url=path, text=hyperlink_string)

    def add_type_caption(self, caption_type, caption_string):
        """Add table caption
        """

        # Add basestring type captions in one step
        if isinstance(caption_string, str):
            # Directly insert base strings
            c = self.add_caption("{} ".format(caption_type), caption_string)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add multi front captions in two steps
        elif isinstance(caption_string, argMultiFontStringHelper):
            # Initiate caption with caption type
            c = self.add_caption("{} ".format(caption_type), "")

            # Insert multi-font string into the report
            self.generate_multi_font_string(
                caption_string, c)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Append Word runs created from multi-font string
            caption_string.execute_backend()

    def add_figure(self, arguments):
        """Add figure to report
        """

        # A figure width is required
        try:
            figure_width = arguments["width"]
        except KeyError:
            print("*  WARNING: no figure width was provided.")
            return

        # Only float numbers of centimeters or inches are understood
        try:
            width_value, width_unit = float(figure_width[:-2]), figure_width[-2:].lower()
        except:
            print("*  WARNING: could not interpret {} as figure width.".format(
                figure_width))
            return
        if width_unit == "mm":
            actual_width = Mm(width_value)
        elif width_unit == "cm":
            actual_width = Cm(width_value)
        elif width_unit == "in":
            actual_width = Inches(width_value)
        else:
            print("*  WARNING: unknown width unit: {}.".format(
                width_unit))
            return

        # Retrieve image and associated caption for figure creation
        figure_file_name, caption_string, from_file = self.fetch_image_and_caption(
            arguments)

        # Create figure when available
        if figure_file_name:
            self.Report.add_picture(
                os.path.join(self.Parameters.OutputDir, figure_file_name),
                width=actual_width)
            p = self.Report.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Append caption when available
            if caption_string:
                self.add_type_caption("Figure", caption_string)

    def add_caption(self, caption_type, caption_string):
        """ Create caption area in XML, with proper number
        """

        # Create new paragraph to populate
        paragraph = self.Report.add_paragraph(caption_type.capitalize().strip() + ' ', style="Caption")
        r = paragraph.add_run()._r

        # Create begin XML element
        beginXml = OxmlElement("w:fldChar")
        beginXml.set(qn("w:fldCharType"), "begin")
        beginXml.set(qn("w:dirty"), "true")

        # Create end XML element
        endXml = OxmlElement("w:fldChar")
        endXml.set(qn("w:fldCharType"), "end")

        # Create number XML element
        numberXml = OxmlElement("w:instrText")
        numberXml.text = "SEQ " + caption_type.upper() + " \\* ARABIC"

        # Add all elements to current run
        r.append(beginXml)
        r.append(numberXml)
        r.append(endXml)

        # Add caption string run
        paragraph.add_run(': ' + caption_string)

        return paragraph

    def recursively_build_report(self, item_tree, list_lvl=1, p=None):
        """Recursively traverse structure tree to build Word report
        """

        # Iterate over each dictionary in the array
        for item in item_tree:
            item_type = item.get('n')
            if item_type == "string":
                # Directly insert text fragment into the report
                self.Report.append(NoEscape(item.get("string")))

            # Handle chapter case
            elif item_type.startswith("chapter"):
                # Create chapter
                self.add_chapter(item)

            # Handle section case
            elif item_type.startswith("section"):
                # Create section
                self.add_section(item)

            # Handle subsection case
            elif item_type.startswith("subsection"):
                # Create subsection
                self.add_subsection(item)

            # Handle subsubsection case
            elif item_type.startswith("subsubsection"):
                # Create subsubsection
                self.add_subsubsection(item)

            # Handle paragraph case
            elif item_type == "paragraph":
                # Add paragraph to the report
                self.add_paragraph(item)

            # Handle itemization/enumeration case
            elif item_type in ("itemize", "enumerate"):
                # Add itemization/enumeration to the report
                self.add_list(item, list_lvl, True if item_type == "enumerate" else False)
                list_lvl = list_lvl + 1

            # Handle verbatim case
            elif item_type == "verbatim":
                # Include file verbatim as report paragraph
                file_name = os.path.join(self.Parameters.DataDir,
                                         item["include"])
                verbatim_style = self.Report.styles["macro"]
                with open(file_name, 'r') as tex_file:

                    # Read each row and add it to Word output
                    for line in tex_file:
                        self.Report.add_paragraph(line).style = verbatim_style

            # Handle figure case
            elif item_type == "figure":
                # Add figure to the report
                self.add_figure(item["arguments"])

            # Handle meta-information case
            elif item_type == "meta":
                # Append meta-information to the report
                self.add_meta_information(item)

            # Handle properties case
            elif item_type == "properties":
                # Append information specific to given property to the report
                self.add_information(item)

            # Handle aggregate case
            elif item_type == "aggregate":
                # Append aggregated information
                self.add_aggregation(item)

            # Handle color-table
            elif item_type == "color-table":
                # Append color-table
                self.add_color_table(data=item)

            # Handle inline document
            elif item_type == "inline-docx":
                # Append inline-docx
                self.inline_docx(data=item)

            # Handle html
            elif item_type == "html":
                # Append html
                self.add_html(data=item)

            # Proceed with recursion if needed
            if "sections" in item:
                self.recursively_build_report(item["sections"], list_lvl)

    def assemble(self, report_map, version=None, latex_proc=None):
        """Create a Word report from given report map
        """

        # Call super method first
        super().assemble(report_map, version)

        # Create report with preamble
        print("[argWordBackend] Generating Word report")
        self.create_document_preamble(version)

        # Pass over data for report contents
        self.recursively_build_report(report_map["chapters"], list_lvl=1)

        # Append backend-specific postamble
        self.append_document_postamble()

        # Handle title numbering
        # self.add_header_numbering()

        # Create Word report
        print("[argWordBackend] Saving Word report")
        self.Report.save("{}.docx".format(os.path.join(
            self.Parameters.OutputDir,
            self.Parameters.FileName)))

        # Update tables of contents, figures, tables on Windows only for now
        if self.osn == "nt":
            self.update_tocs(r"{}.docx".format(os.path.abspath(os.path.join(
                self.Parameters.OutputDir,
                self.Parameters.FileName))))
        else:
            print("*  WARNING: Table of contents not automatically generated yet by Word backend on {} systems.".format(
                self.osn))

    def add_document_provenance(self, version=None):
        """Add document provenance information on a new page
        """

        # Start new page
        self.Report.add_page_break()

        # Compute provenance information
        timestamp = r'This document was generated on {} with the Automatic Report Generator (ARG) version "{}" on the {} system {}.'.format(
            self.get_timestamp(),
            version,
            platform.system(),
            platform.node())
        self.Report.add_paragraph(timestamp)

    def add_document_tocs(self):
        """Add document Table of Contents, List of Figures and List of Tables on new pages
        """

        # Start new page
        self.Report.add_page_break()

        # Add tables of contents, figures, tables
        self.add_table_of_contents()
        self.Report.add_page_break()
        self.add_index("Figure")
        self.Report.add_page_break()
        self.add_index("Table")
        self.Report.add_page_break()

    @staticmethod
    def update_tocs(docx_file):
        """Update table of contents, list of figures, list of tables
        """

        # Try opening Word Application
        try:
            word = Word.Application()
        # Find work-around
        except:
            # Try to retrieve MS Word application
            try:
                word = DispatchEx("Word.Application")
            except:
                # Try running C#-built exe solution
                try:
                    subprocess.check_call([r"../../arg/Backend/updateWordFields", docx_file])
                    return

                # Return with warning message
                except:
                    print("*  WARNING: Word report cannot be automatically updated.")
                    return

        # Open the docx file in MS Word
        doc = word.Documents.Open(docx_file)
        word.ActiveDocument.Fields.Update()

        # Update all tables
        if doc.TablesOfContents(1):
            doc.TablesOfContents(1).Update()
        if doc.TablesOfFigures(1):
            doc.TablesOfFigures(1).Update()
        if doc.TablesOfFigures(2):
            doc.TablesOfFigures(2).Update()

        # Close file and exit cleanly
        doc.Close(SaveChanges=True)
        word.Quit()

    def add_table_of_contents(self, heading_run=None, body_run=None, lvls="1-4"):
        """Add table of contents
        """

        # Add a heading in existing run if provided
        if heading_run:
            heading_run.text = "Table of Contents"
        # Otherwise, create a heading
        else:
            self.Report.add_heading("Table of Contents")

        # Use body run if provided
        if body_run:
            run = body_run

        # Otherwise, create a new run
        else:
            paragraph = self.Report.add_paragraph()
            run = paragraph.add_run()

        # Add a beginning character
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        # Create an interpreted XML text element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')

        # Add level 1 to level 3 contents
        instrText.text = "TOC \\o {} \\h \\z \\u".format(lvls)  # change 1-4 depending on heading levels you need

        # Add a separate character
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        # Add an instruction XMl text element to above character
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        # And an ending character
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        # Add all above XML elements to current run
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)

    def add_index(self, string, heading_run=None, body_run=None):
        """Add table of contents
        """

        heading = "List of " + string.capitalize().strip() + "s"

        # Add a heading in existing run if provided
        if heading_run:
            heading_run.text = heading

        # Otherwise, create a heading
        else:
            self.Report.add_heading("List of " + string.capitalize().strip() + "s")

        # Use body run if provided
        if body_run:
            run = body_run

        # Otherwise, create a new run
        else:
            paragraph = self.Report.add_paragraph()
            run = paragraph.add_run()

        # Add a beginning character
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        # Create an interpreted XML text element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')

        # Add report elements defined by string
        instrText.text = 'TOC \\h \\z \\c "' + string + '" CHARFORMAT'

        # Add an update XML element
        fldChar2 = OxmlElement('w:updateFields')

        # Add an ending character
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        # Add all avobe XML elements to current run
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(fldChar2)
        r_element.append(instrText)
        r_element.append(fldChar4)
        p_element = paragraph._p

    @staticmethod
    def iter_heading(paragraphs):
        """Iterate on headings to number them, a fix suggested by
           https://github.com/python-openxml/python-docx/issues/590
        """

        # Iterate over paragraphs
        for paragraph in paragraphs:
            isItHeading = re.match('Heading ([1-9])', paragraph.style.name)
            # Issue an index only when a heading is encountered
            if isItHeading:
                yield int(isItHeading.groups()[0]), paragraph

    def add_header_numbering(self):
        """Create the sub-tiered header number
        """

        # Allow for up to a maximum of five tiers
        hNums = [0, 0, 0, 0, 0]

        # Iterate over headings
        for index, hx in self.iter_heading(self.Report.paragraphs):
            # Put zeroes in tiers below current one
            for i in range(index + 1, 5):
                hNums[i] = 0

            # Increment current level
            hNums[index] += 1

            # Prepare tiered header index string
            hStr = ''
            for i in range(1, index + 1):
                hStr += "%d." % hNums[i]

            # Add the numbering
            hx.text = "{} {}".format(
                hStr,
                hx.text)

    def parse_colors(self, headers: list, rows: list, appl: str = 'background') -> list:
        """ Parses colors for whole table
        """
        colr = {'background': 1, 'foreground': 2}
        ct = colr.get(appl)

        headers_colors = [header[ct] for header in headers]
        rows_colors = [[cell[ct] for cell in row] for row in rows]
        table_colors = list()
        table_colors.append(headers_colors)
        table_colors.extend(rows_colors)
        table_colors_hex = [[self.parse_rgb_to_hex(color, ct) for color in color_row] for color_row in table_colors]
        return table_colors_hex

    def support_string(self, item: dict) -> Paragraph:
        """ Function created in order not to DRY
        """
        # Retrieve provided string
        string = item["string"]

        # Distinguish between plain and multi-font strings
        if isinstance(string, argMultiFontStringHelper):
            # Insert string into the report
            return self.generate_multi_font_string(string, self.Report.add_paragraph())
        else:
            # Directly insert Word fragment into the report
            return self.Report.add_paragraph(self.apply_break_lines(no_lb_string=string))

    @staticmethod
    def apply_break_lines(no_lb_string: str) -> str:
        """ Takes string without break lines"""
        lb_string = no_lb_string.replace('\\', '\n')
        return lb_string

    @staticmethod
    def parse_vertical_alignment(headers: list, rows: list) -> list:
        """ Parses vertical alignment for whole table
        """
        va_dict = {'c': 'center', 't': 'top', 'b': 'bottom'}
        headers_va = [header[4] for header in headers]
        rows_va = [[cell[4] for cell in row] for row in rows]
        v_a = list()
        v_a.append(headers_va)
        v_a.extend(rows_va)
        table_va = [[va_dict.get(cell_va, 'center') for cell_va in va_row] for va_row in v_a]
        return table_va

    @staticmethod
    def parse_rgb_to_hex(rgb: str, appl: int) -> str:
        """ Parses RGB color to HEX
        """
        if rgb == '' and appl == 1:
            return 'FFFFFF'
        elif rgb == '' and appl == 2:
            return '000000'
        else:
            hex_str = ''.join(
                [hex(int(color))[2:].upper() if len(hex(int(color))[2:]) > 1 else f"0{hex(int(color))[2:].upper()}"
                 for color in rgb.split(',')])
            return hex_str

    @staticmethod
    def parse_cell_align(align: str) -> str:
        """ Parses alignment letter (c, l ,r) to paragrapth alignment
        """
        al_dict = {'c': WD_ALIGN_PARAGRAPH.CENTER, 'r': WD_ALIGN_PARAGRAPH.RIGHT, 'l': WD_ALIGN_PARAGRAPH.LEFT}
        return al_dict.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    @staticmethod
    def add_hyperlink_to_paragraph(paragraph, url, text, color=None, underline=None):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :param color: The color displayed for the url
        :param underline: The underline for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Add color if it is given
        if color is not None:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)

        # Remove underlining if it is requested
        if underline is not None:
            u = docx.oxml.shared.OxmlElement('w:u')
            u.set(docx.oxml.shared.qn('w:val'), 'single')
            rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink, paragraph

    @staticmethod
    def add_string_to_paragraph(paragraph, text) -> Paragraph:
        """
        A function that places a hyperlink within a paragraph object.
        :param paragraph: The paragraph we are adding the hyperlink to.
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        paragraph._p.append(new_run)

        return paragraph
